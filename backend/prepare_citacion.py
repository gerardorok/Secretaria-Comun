"""
Script de preparación de las plantillas del módulo CITACIÓN A NOTIFICACIÓN
PERSONAL.

Procesa los dos archivos institucionales originales:
    - FORMATO_CITACION_PERSONAL.docx              (citación base)
    - FORMATO_CITACION_PERSONAL_VINCULACION.docx  (citación con vinculación)

Y produce las versiones con placeholders Jinja2:
    - plantilla_citacion_base.docx
    - plantilla_citacion_vinculacion.docx

CORREO opcional: el párrafo del correo se envuelve con `{%p if CORREO %}...
{%p endif %}` para que docxtpl elimine el párrafo completo cuando el correo
viene vacío (sin dejar líneas en blanco).

Ejecutar:
    cd /app/backend && python prepare_citacion.py
"""
from pathlib import Path

from docx import Document

ROOT = Path(__file__).parent
SOURCES = {
    "base": ROOT / "templates" / "FORMATO_CITACION_PERSONAL.docx",
    "vinculacion": ROOT / "templates" / "FORMATO_CITACION_PERSONAL_VINCULACION.docx",
}
TARGETS = {
    "base": ROOT / "templates" / "plantilla_citacion_base.docx",
    "vinculacion": ROOT / "templates" / "plantilla_citacion_vinculacion.docx",
}


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    full_text = "".join(r.text for r in paragraph.runs)
    if old not in full_text:
        return False
    start = full_text.find(old)
    end = start + len(old)
    pos = 0
    first_match_done = False
    for run in paragraph.runs:
        run_start = pos
        run_end = pos + len(run.text)
        pos = run_end
        if run_end <= start or run_start >= end:
            continue
        if not first_match_done:
            prefix = run.text[: max(0, start - run_start)]
            suffix = ""
            if run_end > end:
                suffix = run.text[end - run_start :]
            run.text = prefix + new + suffix
            first_match_done = True
        else:
            if run_end > end:
                run.text = run.text[end - run_start :]
            else:
                run.text = ""
    return True


def replace_once(paragraph, old: str, new: str) -> bool:
    if old in new:
        return replace_in_runs(paragraph, old, new)
    found = False
    count = 0
    while replace_in_runs(paragraph, old, new):
        found = True
        count += 1
        if count > 20:
            break
    return found


# Mapeo por índice de párrafo común a ambas plantillas
COMMON_BY_INDEX = {
    2: [
        ("Santa Marta D.T.C.H., \t", "Santa Marta D.T.C.H., {{FECHA_ELABORACION}}"),
        ("Santa Marta D.T.C.H., ", "Santa Marta D.T.C.H., {{FECHA_ELABORACION}}"),
        ("Santa Marta D.T.C.H.", "Santa Marta D.T.C.H., {{FECHA_ELABORACION}}"),
    ],
    # 5: "Señor" -> "Señor(a)"
    5: [("Señor", "Señor(a)")],
}

# Mapeo por plantilla: nombre, dirección, ciudad, correo (índice -> reemplazo)
SPECIFIC = {
    "base": {
        6: ("CRISTIANO RONALDO DOS SANTOS", "{{NOMBRE}}"),
        7: ("Carrera 12 No. 15 – 64", "{{DIRECCION}}"),
        8: ("Santa Marta, Magdalena", "{{CIUDAD}}"),
    },
    "vinculacion": {
        6: ("CRISTIANO RONALDO DOS SANTOS", "{{NOMBRE}}"),
        7: ("Calle 34 No. 33 – 27", "{{DIRECCION}}"),
        8: ("Santa Marta, Magdalena", "{{CIUDAD}}"),
    },
}


def replace_correo_paragraph(doc) -> None:
    """Reemplaza el contenido del párrafo del correo por el placeholder
    `{{ CORREO }}` y lo envuelve con párrafos de control Jinja2
    `{%p if CORREO %}` / `{%p endif %}` para que docxtpl elimine el
    párrafo completo cuando el correo viene vacío.

    En la plantilla original el correo está dentro de un `<w:hyperlink>`,
    así que se manipula el XML directamente:
      1. Localiza el párrafo 9 (correo).
      2. Conserva el formato (rPr) del run del hipervínculo.
      3. Borra el hyperlink y deja un único run con `{{ CORREO }}`.
      4. Inserta antes y después del párrafo dos párrafos de control con
         las directivas `{%p if CORREO %}` y `{%p endif %}` (que docxtpl
         remueve siempre, quedando o no el correo según el contexto).
    """
    from copy import deepcopy
    from docx.oxml.ns import qn

    par9 = doc.paragraphs[9]
    p = par9._p

    # 1. Capturar formato del run del hyperlink
    rpr_template = None
    for hyper in p.findall(qn("w:hyperlink")):
        for r in hyper.findall(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is not None:
                rpr_template = deepcopy(rpr)
                break
        if rpr_template is not None:
            break

    # 2. Limpiar el párrafo (excepto pPr)
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)

    # 3. Agregar nuevo run con {{ CORREO }}
    new_r = p.makeelement(qn("w:r"), {})
    if rpr_template is not None:
        new_r.append(rpr_template)
    t = p.makeelement(qn("w:t"), {})
    t.set(qn("xml:space"), "preserve")
    t.text = "{{CORREO}}"
    new_r.append(t)
    p.append(new_r)

    # 4. Insertar párrafos de control antes y después
    def make_control_paragraph(text: str):
        new_p = p.makeelement(qn("w:p"), {})
        # Copiamos pPr del párrafo del correo para mantener alineación
        ppr_src = p.find(qn("w:pPr"))
        if ppr_src is not None:
            new_p.append(deepcopy(ppr_src))
        r = new_p.makeelement(qn("w:r"), {})
        t = new_p.makeelement(qn("w:t"), {})
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        new_p.append(r)
        return new_p

    p_if = make_control_paragraph("{%p if CORREO %}")
    p_end = make_control_paragraph("{%p endif %}")

    p.addprevious(p_if)
    p.addnext(p_end)

# Reemplazos del cuerpo (párrafo 17) - BASE
BODY_BASE = [
    (
        "POR MEDIO DEL CUAL SE ORDENA LA APERTURA DEL PROCESO DE RESPONSABILIDAD FISCAL",
        "{{NOMBRE_PROVIDENCIA}}",
    ),
    (
        "Gerencia Departamental Colegiada de Magdalena de la Contraloría General de la Republica",
        "{{PROFERIDO_POR}}",
    ),
    ("MUNICIPIO DE PLATO", "{{ENTIDAD_AFECTADA}}"),
    ("AUTO No. 419", "{{TIPO_PROVIDENCIA}} No. {{NUMERO_PROVIDENCIA}}"),
    ("23 de octubre de 2025", "{{FECHA_PROVIDENCIA}}"),
    ("PRF-80472-2024-46137", "PRF-{{PRF}}"),
]

# Reemplazos del cuerpo (párrafo 17) - VINCULACIÓN
# Orden importa: providencia 2 primero (es la posterior en el texto)
BODY_VINCULACION = [
    (
        "POR MEDIO DEL CUAL SE ORDENA LA APERTURA E IMPUTACIÓN DENTRO DEL PROCESO DE RESPONSABILIDAD FISCAL",
        "{{NOMBRE_PROVIDENCIA}}",
    ),
    (
        "POR MEDIO DEL CUAL SE VINCULA E IMPUTA EN EL PROCESO DE RESPONSABILIDAD FISCAL",
        "{{NOMBRE_PROVIDENCIA_2}}",
    ),
    (
        "GERENCIA DEPARTAMENTAL COLEGIADA DE MAGDALENA DE LA CONTRALORÍA GENERAL DE LA REPUBLICA",
        "{{PROFERIDO_POR}}",
    ),
    ("MUNICIPIO DE EL PIÑON", "{{ENTIDAD_AFECTADA}}"),
    ("AUTO No. 420", "{{TIPO_PROVIDENCIA}} No. {{NUMERO_PROVIDENCIA}}"),
    ("AUTO No. 013", "{{TIPO_PROVIDENCIA_2}} No. {{NUMERO_PROVIDENCIA_2}}"),
    ("23 de octubre de 2025", "{{FECHA_PROVIDENCIA}}"),
    ("22 de enero de 2026", "{{FECHA_PROVIDENCIA_2}}"),
    ("PRF-80472-2025-49159", "PRF-{{PRF}}"),
]

BODY = {"base": BODY_BASE, "vinculacion": BODY_VINCULACION}


def process(kind: str) -> None:
    source = SOURCES[kind]
    target = TARGETS[kind]
    if not source.exists():
        raise FileNotFoundError(f"Falta la plantilla original: {source}")

    doc = Document(str(source))

    # 1. Reemplazos comunes por índice
    for idx, replacements in COMMON_BY_INDEX.items():
        if idx < len(doc.paragraphs):
            par = doc.paragraphs[idx]
            for old, new in replacements:
                # Idempotencia: si el placeholder destino ya está, saltamos
                if new.startswith("Santa Marta") and "{{FECHA_ELABORACION}}" in par.text:
                    break
                if new == "Señor(a)" and "Señor(a)" in par.text:
                    break
                if replace_once(par, old, new):
                    break

    # 2. Reemplazos específicos por plantilla
    for idx, (old, new) in SPECIFIC[kind].items():
        if idx < len(doc.paragraphs):
            par = doc.paragraphs[idx]
            if new in par.text:
                continue
            replace_once(par, old, new)

    # 2b. Párrafo del correo (manipulación XML por el hipervínculo)
    replace_correo_paragraph(doc)

    # 3. Reemplazos en el cuerpo
    for par in doc.paragraphs:
        for old, new in BODY[kind]:
            replace_once(par, old, new)

    doc.save(str(target))

    # Verificación
    out = Document(str(target))
    text = "\n".join(p.text for p in out.paragraphs)
    expected = [
        "{{NOMBRE}}",
        "{{DIRECCION}}",
        "{{CIUDAD}}",
        "{{CORREO}}",
        "{{TIPO_PROVIDENCIA}}",
        "{{NUMERO_PROVIDENCIA}}",
        "{{FECHA_PROVIDENCIA}}",
        "{{NOMBRE_PROVIDENCIA}}",
        "{{PRF}}",
        "{{ENTIDAD_AFECTADA}}",
        "{{PROFERIDO_POR}}",
        "{{FECHA_ELABORACION}}",
    ]
    if kind == "vinculacion":
        expected += [
            "{{TIPO_PROVIDENCIA_2}}",
            "{{NUMERO_PROVIDENCIA_2}}",
            "{{FECHA_PROVIDENCIA_2}}",
            "{{NOMBRE_PROVIDENCIA_2}}",
        ]
    print(f"\n=== {kind.upper()} -> {target.name} ===")
    for v in expected:
        present = "OK" if v in text else "FALTA"
        print(f"  {present:6} {v}")


def main() -> None:
    for kind in ("base", "vinculacion"):
        process(kind)


if __name__ == "__main__":
    main()
