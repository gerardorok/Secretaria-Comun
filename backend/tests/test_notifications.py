"""Tests for the CGR Notification Generator backend.

Covers:
  - GET /api/                       (health check)
  - GET /api/template/status        (template availability)
  - POST /api/generate-notification (full happy-path with .docx parsing)
  - POST /api/generate-notification (validation / 422)
  - ENTIDAD reuse rule (page-2 ENTIDAD == page-1 ENTIDAD_AFECTADA)
  - Document structural integrity (1 table with 16 rows)
"""

import io
import os
import re

import pytest
import requests
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "https://aviso-generador.preview.emergentagent.com").rstrip("/")
API = f"{BASE_URL}/api"


# ----------------------------- Fixtures -------------------------------------
@pytest.fixture(scope="module")
def api_client():
    s = requests.Session()
    s.headers.update({"Content-Type": "application/json"})
    return s


@pytest.fixture
def valid_payload():
    return {
        "NOMBRE": "TEST_RODRIGO ALBERTO ANDRADE ORTIZ",
        "DIRECCION": "Diagonal 4 a 2 a - 22, Barrio Loma Fresca",
        "CIUDAD": "Arigani, Magdalena",
        "NUMERO_AVISO": "029-2026",
        "AUTO": "AUTO No. 503",
        "FECHA_AUTO": "27 de noviembre de 2025",
        "PRF": "80472-2025-48809",
        "ENTIDAD_AFECTADA": "MUNICIPIO DE SANTA ANA",
        "PROFERIDO_POR": "GERENCIA DEPARTAMENTAL COLEGIADA DE MAGDALENA",
        "NOTIFICADO": "JUAN PEREZ GOMEZ",
        "FECHA_ELABORACION": "28 de noviembre de 2025",
        "PERSONAS_NOTIFICAR": "JUAN PEREZ, MARIA LOPEZ",
        "PROVIDENCIA": "AUTO No. 503",
        "FECHA_PROVIDENCIA": "27 de noviembre de 2025",
        "TIPO_PROVIDENCIA": "AUTO DE APERTURA",
        "FECHA_CITACION": "01/12/2025",
        "ANEXO": "AUTO No. 503 en 19 folios",
    }


# ----------------------------- Health / Status ------------------------------
class TestHealth:
    def test_root_health(self, api_client):
        r = api_client.get(f"{API}/")
        assert r.status_code == 200
        body = r.json()
        assert "message" in body
        assert "CGR" in body["message"] or "Aviso" in body["message"]

    def test_template_status(self, api_client):
        r = api_client.get(f"{API}/template/status")
        assert r.status_code == 200
        body = r.json()
        assert body.get("exists") is True
        assert "plantilla_aviso.docx" in body.get("template_path", "")
        assert body.get("size_bytes", 0) > 0


# ----------------------------- Validation -----------------------------------
class TestValidation:
    def test_missing_required_field_returns_422(self, api_client, valid_payload):
        payload = dict(valid_payload)
        payload.pop("NOMBRE")
        r = api_client.post(f"{API}/generate-notification", json=payload)
        assert r.status_code == 422

    def test_empty_string_required_field_returns_422(self, api_client, valid_payload):
        payload = dict(valid_payload)
        payload["NOMBRE"] = ""
        r = api_client.post(f"{API}/generate-notification", json=payload)
        assert r.status_code == 422

    def test_completely_empty_body_returns_422(self, api_client):
        r = api_client.post(f"{API}/generate-notification", json={})
        assert r.status_code == 422


# ----------------------------- Document generation --------------------------
class TestGeneration:
    def _generate(self, api_client, payload):
        r = api_client.post(f"{API}/generate-notification", json=payload)
        return r

    def test_generate_returns_docx_with_correct_headers(self, api_client, valid_payload):
        r = self._generate(api_client, valid_payload)
        assert r.status_code == 200, f"body: {r.text[:300]}"
        ct = r.headers.get("content-type", "")
        assert "officedocument.wordprocessingml.document" in ct
        cd = r.headers.get("content-disposition", "")
        assert "attachment" in cd
        m = re.search(r'filename="?([^";]+)"?', cd)
        assert m is not None, f"No filename in content-disposition: {cd}"
        assert m.group(1).lower().endswith(".docx")
        assert len(r.content) > 1000  # actual docx zip

    def test_generated_docx_contains_all_values(self, api_client, valid_payload):
        r = self._generate(api_client, valid_payload)
        assert r.status_code == 200
        doc = Document(io.BytesIO(r.content))

        # Collect all text from paragraphs + tables
        full_text_parts = []
        for p in doc.paragraphs:
            full_text_parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text_parts.append(cell.text)
        full_text = "\n".join(full_text_parts)

        # All input values must appear somewhere
        keys_to_check = [
            "NOMBRE", "DIRECCION", "CIUDAD", "AUTO", "FECHA_AUTO", "PRF",
            "ENTIDAD_AFECTADA", "PROFERIDO_POR", "NOTIFICADO",
            "FECHA_ELABORACION", "PERSONAS_NOTIFICAR", "PROVIDENCIA",
            "FECHA_PROVIDENCIA", "TIPO_PROVIDENCIA", "FECHA_CITACION", "ANEXO",
        ]
        missing = [k for k in keys_to_check if valid_payload[k] not in full_text]
        assert not missing, f"Values not found in generated docx: {missing}"

        # No unrendered Jinja placeholders should remain
        unrendered = re.findall(r"\{\{[^}]+\}\}", full_text)
        assert not unrendered, f"Unrendered Jinja placeholders: {unrendered}"

    def test_entidad_reuse_in_page2(self, api_client, valid_payload):
        """ENTIDAD on page 2 must be auto-filled from ENTIDAD_AFECTADA."""
        r = self._generate(api_client, valid_payload)
        assert r.status_code == 200
        doc = Document(io.BytesIO(r.content))

        target = valid_payload["ENTIDAD_AFECTADA"]
        # Look in tables (page 2 cuadro de notificacion)
        found_in_table = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target in cell.text:
                        found_in_table = True
                        break
        assert found_in_table, f"ENTIDAD_AFECTADA '{target}' not found in any table cell (page 2 reuse failed)"

    def test_document_structure_preserved(self, api_client, valid_payload):
        """Generated doc must keep at least 1 table with 16 rows (cuadro de notificación)."""
        r = self._generate(api_client, valid_payload)
        assert r.status_code == 200
        doc = Document(io.BytesIO(r.content))

        assert len(doc.tables) >= 1, "No table found in generated document"
        # Find a table that has 16 rows (cuadro de notificación page 2)
        row_counts = [len(t.rows) for t in doc.tables]
        assert 16 in row_counts, f"Expected a table with 16 rows; got tables with rows={row_counts}"

        # Should preserve a non-trivial number of paragraphs (page 1 content)
        assert len(doc.paragraphs) > 5

    def test_filename_contains_numero_aviso(self, api_client, valid_payload):
        r = self._generate(api_client, valid_payload)
        assert r.status_code == 200
        cd = r.headers.get("content-disposition", "")
        assert "029" in cd  # NUMERO_AVISO part
