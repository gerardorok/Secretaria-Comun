"""
Script de preparación de la plantilla.

Lee el archivo original AVISO_PRF_FORMATO.docx (que contiene valores de ejemplo
reales y la firma del funcionario embebida) y produce plantilla_aviso.docx
con placeholders Jinja2 ({{ VAR }}) que serán renderizados por docxtpl,
conservando exactamente el formato original (fuentes, tamaños, negrillas,
espaciados, tablas, márgenes, alineación, saltos de página, encabezados,
pies de página y firmas/imágenes).

Ejecutar una sola vez tras cualquier cambio en la plantilla original:
    cd /app/backend && python prepare_template.py
"""
from docx import Document
from pathlib import Path

ROOT = Path(__file__).parent
SOURCE = ROOT / "templates" / "AVISO_PRF_FORMATO.docx"
TARGET = ROOT / "templates" / "plantilla_aviso.docx"


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    """Reemplaza `old` por `new` en un párrafo, posiblemente atravesando
    varios runs, conservando el formato del primer run que coincide.

    Devuelve True si se realizó algún reemplazo.
    """
    full_text = "".join(r.text for r in paragraph.runs)
    if old not in full_text:
        return False

    start = full_text.find(old)
    end = start + len(old)

    pos = 0
    first_match_done = False
    for run in paragraph.runs:
        run_start = pos
        run_end = pos + len(run.text)
        pos = run_end

        if run_end <= start or run_start >= end:
            continue

        if not first_match_done:
            prefix = run.text[: max(0, start - run_start)]
            suffix = ""
            if run_end > end:
                suffix = run.text[end - run_start :]
            run.text = prefix + new + suffix
            first_match_done = True
        else:
            if run_end > end:
                run.text = run.text[end - run_start :]
            else:
                run.text = ""
    return True


def replace_in_paragraph_repeated(paragraph, old: str, new: str) -> int:
    """Reemplaza todas las ocurrencias no superpuestas de `old` por `new`."""
    if old in new:
        return 1 if replace_in_runs(paragraph, old, new) else 0
    count = 0
    while replace_in_runs(paragraph, old, new):
        count += 1
        if count > 20:
            break
    return count


# Reemplazos exclusivos para PÁRRAFOS (página 1 + bloque texto página 2).
# El orden importa: cadenas más largas primero para evitar colisiones.
PARAGRAPH_REPLACEMENTS_GLOBAL = [
    # Párrafo 18 (cuerpo página 1)
    ("AVISO No. 085-2026", "AVISO No. {{NUMERO_AVISO}}"),
    ("AUTO No. 456", "{{AUTO}}"),
    ("5 de mayo de 2026", "{{FECHA_AUTO}}"),
    ("PRF-80472-2025-48999", "PRF-{{PRF}}"),
    (
        "GERENCIA DEPARTAMENTAL COLEGIADA DE MAGDALENA DE LA CONTRALORIA GENERAL DE LA REPUBLICA",
        "{{PROFERIDO_POR}}",
    ),
    ("MUNICIPIO DE TENERIFE", "{{ENTIDAD_AFECTADA}}"),
    # Dirección y ciudad
    ("Calle 78 No. 14 – 15", "{{DIRECCION}}"),
    ("Santa Marta, Magdalena", "{{CIUDAD}}"),
]

# Reemplazos por índice de párrafo (cuando el mismo texto aparece en lugares
# distintos y debe mapearse a variables diferentes).
PARAGRAPH_REPLACEMENTS_BY_INDEX = {
    # Párrafo 4: añadir fecha de elaboración tras la ciudad de elaboración.
    # Manejamos dos formatos posibles según versión de la plantilla:
    4: [
        ("Santa Marta D.T.C.H., ", "Santa Marta D.T.C.H., {{FECHA_ELABORACION}}"),
        ("Santa Marta D.T.C.H.", "Santa Marta D.T.C.H., {{FECHA_ELABORACION}}"),
    ],
    # Párrafo 8: encabezado inclusivo
    8: [("Señor", "Señor(a)")],
    # Párrafo 9: nombre del destinatario (página 1)
    9: [("CRISTIANO RONALDO DOS SANTOS", "{{NOMBRE}}")],
    # Párrafo 43: persona notificada (página 2)
    43: [("CRISTIANO RONALDO DOS SANTOS", "{{NOTIFICADO}}")],
}

# Reemplazos para celdas de TABLAS (página 2). Largos primero.
TABLE_REPLACEMENTS = [
    ("AUTO 456 en 17 folios", "{{ANEXO}}"),
    (
        "CRISTIANO RONALDO DOS SANTOS, LIONEL ANDRES MESSI",
        "{{PERSONAS_NOTIFICAR}}",
    ),
    (
        "GERENCIA DEPARTAMENTAL COLEGIADA DE MAGDALENA DE LA CONTRALORIA GENERAL DE LA REPUBLICA",
        "{{PROFERIDO_POR}}",
    ),
    ("Aviso No.085-2026", "Aviso No.{{NUMERO_AVISO}}"),
    ("7 de mayo de 2026", "{{FECHA_ELABORACION}}"),
    ("AUTO No. 456", "{{PROVIDENCIA}}"),
    ("5 de mayo de 2026", "{{FECHA_PROVIDENCIA}}"),
    ("AUTO DE APERTURA", "{{TIPO_PROVIDENCIA}}"),
    ("80472-2025-48999", "{{PRF}}"),
    ("MUNICIPIO DE TENERIFE", "{{ENTIDAD}}"),
    # Fecha de citación (varía: el ejemplo viene en blanco o como dd/mm/aaaa).
    # Lo dejamos como una segunda pasada por si está presente.
]

# Algunas plantillas pueden tener la fecha de citación en formato dd/mm/aaaa.
# Probamos varios candidatos comunes.
FECHA_CITACION_CANDIDATES = ["19/12/2025", "06/05/2026", "05/05/2026"]


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"No se encontró la plantilla original: {SOURCE}")

    doc = Document(str(SOURCE))

    # 1. Reemplazos por índice de párrafo
    for idx, replacements in PARAGRAPH_REPLACEMENTS_BY_INDEX.items():
        if idx < len(doc.paragraphs):
            par = doc.paragraphs[idx]
            for old, new in replacements:
                # Idempotencia: si el texto destino ya está presente,
                # asumimos que la plantilla ya fue editada manualmente.
                if new.strip("{} ") and new in par.text:
                    continue
                # Solo aplicamos el primer patrón que coincida
                if replace_in_paragraph_repeated(par, old, new):
                    break

    # 2. Reemplazos globales en todos los párrafos
    for par in doc.paragraphs:
        for old, new in PARAGRAPH_REPLACEMENTS_GLOBAL:
            replace_in_paragraph_repeated(par, old, new)

    # 3. Reemplazos en celdas de tablas
    processed_tcs = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                if any(tc is x for x in processed_tcs):
                    continue
                processed_tcs.append(tc)

                for par in cell.paragraphs:
                    for old, new in TABLE_REPLACEMENTS:
                        replace_in_paragraph_repeated(par, old, new)
                    # Fecha de citación: probamos varios formatos posibles
                    for cand in FECHA_CITACION_CANDIDATES:
                        if replace_in_paragraph_repeated(par, cand, "{{FECHA_CITACION}}"):
                            break

    doc.save(str(TARGET))
    print(f"OK - Plantilla generada: {TARGET}")

    # Verificación: contamos placeholders insertados
    out = Document(str(TARGET))
    full = []
    for p in out.paragraphs:
        full.append(p.text)
    for t in out.tables:
        for r in t.rows:
            for c in r.cells:
                full.append(c.text)
    text = "\n".join(full)
    expected = [
        "{{NOMBRE}}",
        "{{DIRECCION}}",
        "{{CIUDAD}}",
        "{{NUMERO_AVISO}}",
        "{{AUTO}}",
        "{{FECHA_AUTO}}",
        "{{PRF}}",
        "{{ENTIDAD_AFECTADA}}",
        "{{PROFERIDO_POR}}",
        "{{NOTIFICADO}}",
        "{{FECHA_ELABORACION}}",
        "{{PERSONAS_NOTIFICAR}}",
        "{{PROVIDENCIA}}",
        "{{FECHA_PROVIDENCIA}}",
        "{{TIPO_PROVIDENCIA}}",
        "{{ENTIDAD}}",
        "{{FECHA_CITACION}}",
        "{{ANEXO}}",
    ]
    print("\nPlaceholders encontrados:")
    for v in expected:
        present = "OK" if v in text else "FALTA"
        print(f"  {present:6} {v}")


if __name__ == "__main__":
    main()
