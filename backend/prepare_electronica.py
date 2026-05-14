"""
Script de preparación de las plantillas del módulo de NOTIFICACIÓN ELECTRÓNICA.

Procesa los dos archivos institucionales originales:
    - FORMATO_ELECTRONICA_NATURAL.docx  (persona natural)
    - FORMATO_ELECTRONICA_JURIDICA.docx (persona jurídica)

Y produce las versiones con placeholders Jinja2:
    - plantilla_electronica_natural.docx
    - plantilla_electronica_juridica.docx

Preserva la firma del funcionario, fuentes, alineaciones y demás formato.

Ejecutar:
    cd /app/backend && python prepare_electronica.py
"""
from pathlib import Path

from docx import Document

ROOT = Path(__file__).parent
SOURCES = {
    "natural": ROOT / "templates" / "FORMATO_ELECTRONICA_NATURAL.docx",
    "juridica": ROOT / "templates" / "FORMATO_ELECTRONICA_JURIDICA.docx",
}
TARGETS = {
    "natural": ROOT / "templates" / "plantilla_electronica_natural.docx",
    "juridica": ROOT / "templates" / "plantilla_electronica_juridica.docx",
}


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    """Reemplaza `old` por `new` en un párrafo, posiblemente atravesando
    varios runs, conservando el formato del primer run que coincide."""
    full_text = "".join(r.text for r in paragraph.runs)
    if old not in full_text:
        return False

    start = full_text.find(old)
    end = start + len(old)

    pos = 0
    first_match_done = False
    for run in paragraph.runs:
        run_start = pos
        run_end = pos + len(run.text)
        pos = run_end

        if run_end <= start or run_start >= end:
            continue

        if not first_match_done:
            prefix = run.text[: max(0, start - run_start)]
            suffix = ""
            if run_end > end:
                suffix = run.text[end - run_start :]
            run.text = prefix + new + suffix
            first_match_done = True
        else:
            if run_end > end:
                run.text = run.text[end - run_start :]
            else:
                run.text = ""
    return True


def replace_once(paragraph, old: str, new: str) -> bool:
    """Aplica una sola sustitución (idempotente)."""
    if old in new:
        return replace_in_runs(paragraph, old, new)
    # Para textos donde `old` NO está en `new` permitimos múltiples ocurrencias
    found_any = False
    count = 0
    while replace_in_runs(paragraph, old, new):
        found_any = True
        count += 1
        if count > 20:
            break
    return found_any


# Reemplazos globales que aplican igual a ambos archivos.
BODY_REPLACEMENTS = [
    # Párrafo 19 (cuerpo): orden importa, hacer los más largos primero
    (
        "POR EL CUAL SE ORDENA EL CIERRE DE LA INDAGACIÓN PRELIMINAR Y LA APERTURA DEL PROCESO DE RESPONSABILIDAD FISCAL",
        "{{NOMBRE_PROVIDENCIA}}",
    ),
    (
        "Gerencia Departamental de Magdalena de la Contraloría General de la República",
        "{{PROFERIDO_POR}}",
    ),
    ("MUNICIPIO DE SABANAS DE SAN ANGEL", "{{ENTIDAD_AFECTADA}}"),
    ("AUTO No. 126", "{{TIPO_PROVIDENCIA}} No. {{NUMERO_PROVIDENCIA}}"),
    ("29 de abril de 2026", "{{FECHA_PROVIDENCIA}}"),
    ("PRF-80472-2024-46647", "PRF-{{PRF}}"),
    # Párrafo 22 (folios)
    ("veintidós (22)", "{{NUMERO_FOLIOS}}"),
]

# Reemplazos específicos por persona (correo y nombre).
PERSONA_SPECIFIC = {
    "natural": {
        11: ("CRISTIANO RONALDO DOS SANTOS", "{{NOMBRE}}"),
        12: ("cr7rmmu@gmail.com", "{{CORREO}}"),
    },
    "juridica": {
        11: ("CARMING SAS", "{{NOMBRE}}"),
        12: ("carmingsas@gmail.com", "{{CORREO}}"),
    },
}

# Párrafo 5: "Santa Marta D.T.C.H., " -> añadir fecha de elaboración auto
PARAGRAPH_5_PATTERNS = [
    ("Santa Marta D.T.C.H., ", "Santa Marta D.T.C.H., {{FECHA_ELABORACION}}"),
    ("Santa Marta D.T.C.H.", "Santa Marta D.T.C.H., {{FECHA_ELABORACION}}"),
]


def process(kind: str) -> None:
    source = SOURCES[kind]
    target = TARGETS[kind]
    if not source.exists():
        raise FileNotFoundError(f"Falta la plantilla original: {source}")

    doc = Document(str(source))

    # 1. Reemplazos específicos por tipo de persona (índice de párrafo)
    for idx, (old, new) in PERSONA_SPECIFIC[kind].items():
        if idx < len(doc.paragraphs) and new not in doc.paragraphs[idx].text:
            replace_once(doc.paragraphs[idx], old, new)

    # 2. Párrafo 5: fecha de elaboración
    if 5 < len(doc.paragraphs):
        par5 = doc.paragraphs[5]
        if "{{FECHA_ELABORACION}}" not in par5.text:
            for old, new in PARAGRAPH_5_PATTERNS:
                if replace_once(par5, old, new):
                    break

    # 3. Reemplazos globales en todos los párrafos
    for par in doc.paragraphs:
        for old, new in BODY_REPLACEMENTS:
            replace_once(par, old, new)

    doc.save(str(target))

    # Verificación
    out = Document(str(target))
    text = "\n".join(p.text for p in out.paragraphs)
    expected = [
        "{{NOMBRE}}",
        "{{CORREO}}",
        "{{TIPO_PROVIDENCIA}}",
        "{{NUMERO_PROVIDENCIA}}",
        "{{FECHA_PROVIDENCIA}}",
        "{{NOMBRE_PROVIDENCIA}}",
        "{{PRF}}",
        "{{ENTIDAD_AFECTADA}}",
        "{{PROFERIDO_POR}}",
        "{{NUMERO_FOLIOS}}",
        "{{FECHA_ELABORACION}}",
    ]
    print(f"\n=== {kind.upper()} -> {target.name} ===")
    for v in expected:
        present = "OK" if v in text else "FALTA"
        print(f"  {present:6} {v}")


def main() -> None:
    for kind in ("natural", "juridica"):
        process(kind)


if __name__ == "__main__":
    main()
